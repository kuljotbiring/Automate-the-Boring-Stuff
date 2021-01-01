import docx, os

os.chdir('/Users/kj/Desktop/Certificates/Automate the Boring Stuff')

# return a list of document objects
d = docx.Document('demo.docx')

# get the paragraph objects and the string in the first one
p = d.paragraphs[0].txt

# each change in style/run
run = p.runs[0].text

# change the formatting of the run
p.runs[0].bold = None
p.runs[0].italic = True
p.runs[0].underline = True
p.runs[0].text = 'italic and underlined'
p.style = 'Title'


# save the document
d.save('demo2.docx')

# create a document and add paragraphs
d = docx.Document()
d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('Hello this is another paragraph')

p = d.paragraphs[0]
p.runs[0].bold = True

# save the document
d.save('demo3.docx')

def getText(fileName):
    doc = docx.Document(fileName)
    fullText = []
    for paragraph in doc.paragraphs:
        fullText.append(paragraph.text)
        '\n'.join(fullText)

print(getText('demo.docx'))
